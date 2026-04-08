from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

# Styles helper
def set_font(run, name='Calibri', size=10, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E74B5')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    set_font(run, size=10.5, bold=True, color=(46, 116, 181))
    add_hr(doc)

def add_para(doc, space_before=0, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p

# ── HEADER ────────────────────────────────────────────────────
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(2)
r = name_p.add_run('Muhammad Fahad Imdad')
set_font(r, size=18, bold=True, color=(46, 116, 181))

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_after = Pt(6)
r = contact_p.add_run('+92 314 7800991   |   fahadimdad966@gmail.com   |   linkedin.com/in/muhammadfahadimdad   |   fahadimdad.com')
set_font(r, size=9, color=(68, 68, 68))

# ── EDUCATION ─────────────────────────────────────────────────
add_section_heading(doc, 'Education')

p = add_para(doc, space_after=1)
r1 = p.add_run('Salim Habib University')
set_font(r1, bold=True)
r2 = p.add_run('   Karachi, Pakistan')
set_font(r2, color=(100,100,100))

p = add_para(doc, space_after=1)
r1 = p.add_run('B.S. in Computer Science – CGPA 3.68/4.0   |   Gold Medalist')
set_font(r1)
r2 = p.add_run('   March 2021 – May 2025')
set_font(r2, color=(100,100,100))

p = add_para(doc, space_after=4)
r1 = p.add_run('Private')
set_font(r1, bold=True)
r2 = p.add_run('   Karachi, Pakistan')
set_font(r2, color=(100,100,100))
r3 = p.add_run('   |   GCE AS & A Level')
set_font(r3)
r4 = p.add_run('   July 2019 – Nov 2020')
set_font(r4, color=(100,100,100))

# ── PROJECTS ──────────────────────────────────────────────────
add_section_heading(doc, 'Projects')

projects = [
    ('Robotic Arm via LLMs',
     'Speech-controlled robotic arm simulator using LLMs, speech recognition, and forward kinematics for real-time 3D motion.'),
    ('Robotic Arm via Computer Vision',
     'Hand-tracking robotic arm using MediaPipe and OpenCV with gesture-based control and real-time visualization.'),
    ('AR Circuit Simulator',
     'AR-based electronic circuit simulator using Unity and ARCore; integrated AI model for hardware detection and virtual component overlay.'),
    ('AI-Powered Conversational Shopping Platform',
     'Full-stack e-commerce platform using React, FastAPI and Azure OpenAI with multi-agentic LLMs, LangChain, vector-based RAG, and context-aware recommendations.'),
    ('Medical Claim App',
     'AI-powered app using Azure OCR and Azure OpenAI to extract and validate medical prescriptions against company health policy.'),
    ('Brain Tumor Classification',
     'Deep learning model to classify brain tumors from MRI scans. Presented at Sindh HEC Research Showcase 2024.'),
    ('Visual Defect Detection',
     'ResNet50-based classifier trained on MVTec AD dataset with a Streamlit UI for industrial defect recognition.'),
    ('Vehicle Counting & Speed Detection (YOLOv8 & YOLOv11)',
     'Lane-wise vehicle counting (YOLOv8) and speed estimation (YOLOv11) with real-time vehicle tracking.'),
    ('Fire Detection (YOLOv8)',
     'YOLOv8-based fire detection system trained on video footage for safety and incident analysis.'),
    ('AI Medical Interview Platform',
     'Personalized interview simulator for radiology professionals using Google Gemini, voice-based Q&A with real-time transcription, rubric-based AI feedback, and computer vision proctoring for exam integrity.'),
]

for title, desc in projects:
    p = add_para(doc, space_before=3, space_after=1)
    r1 = p.add_run(title)
    set_font(r1, bold=True)
    p2 = add_para(doc, space_after=3)
    p2.paragraph_format.left_indent = Cm(0.5)
    r2 = p2.add_run(desc)
    set_font(r2, size=9.5)

# ── EXPERIENCE ────────────────────────────────────────────────
add_section_heading(doc, 'Experience')

def add_experience(doc, org, location, role, dates, bullets):
    p = add_para(doc, space_before=4, space_after=1)
    r1 = p.add_run(org)
    set_font(r1, bold=True)
    r2 = p.add_run(f'   {location}')
    set_font(r2, color=(100,100,100))

    p2 = add_para(doc, space_after=2)
    r3 = p2.add_run(role)
    set_font(r3, bold=True, size=9.5)
    r4 = p2.add_run(f'   {dates}')
    set_font(r4, size=9.5, color=(100,100,100))

    for b in bullets:
        p3 = add_para(doc, space_after=1)
        p3.paragraph_format.left_indent = Cm(0.5)
        r5 = p3.add_run(f'• {b}')
        set_font(r5, size=9.5)

add_experience(doc, 'Beam AI', 'Karachi, Pakistan', 'AI Agent Engineer', 'Nov 2025 – Present', [
    'Designed and deployed autonomous AI agents replacing manual business processes with intelligent, fully automated systems.',
    'Built multi-step agent architectures using graph-based workflow systems, LLM orchestration, and external tool integration for enterprise clients including Booth & Partner and UNiDAYS.',
    'Implemented dynamic task routing, error handling, retry logic, and context management for long-running multi-step workflows.',
    'Integrated enterprise systems (Airtable, Slack, email) and developed evaluation frameworks with feedback loops for continuous improvement.',
])

add_experience(doc, 'Systems Limited', 'Karachi, Pakistan', 'AI & Data Science Intern', 'June 2025 – Oct 2025', [
    'Built AI solutions across Computer Vision, NLP, and Generative AI for enterprise use cases.',
    'Led development of Medical Claim App (Azure OCR + Azure OpenAI) and AI-powered chat shopping platform (React, Azure OpenAI).',
    'Trained and evaluated models using Python, Scikit-learn, TensorFlow and Keras; hands-on experience with Azure Machine Learning and Azure Cognitive Services.',
])

add_experience(doc, 'IEEE Computer Society, Salim Habib University', 'Karachi, Pakistan', 'Chairperson', 'July 2024 – Dec 2024', [
    'Led the chapter organising workshops, competitions, and networking events; expanded membership and industry-academia collaborations.',
])

add_experience(doc, 'IEEE SHU Student Branch, Salim Habib University', 'Karachi, Pakistan', 'Head of Operating Committees & Webmaster', 'Dec 2023 – Dec 2024', [
    'Organised TechNexus flagship event; developed Women in Engineering (WIE) Society, COMSOC, and EMBS within IEEE SHU.',
])

add_experience(doc, 'Smart City Lab, NCAI-NEDUET', 'Karachi, Pakistan', 'Research Intern', 'Mar 2023 – Aug 2023', [
    'Developed AR application for electronic circuit simulation using Unity, ARCore and AI; integrated AI model for hardware detection and virtual overlay superimposition.',
])

# ── LEADERSHIP ────────────────────────────────────────────────
add_section_heading(doc, 'Leadership Activities')

p = add_para(doc, space_before=4, space_after=1)
r = p.add_run('NASA Space App Challenge   ')
set_font(r, bold=True)
r2 = p.add_run('Karachi, Pakistan')
set_font(r2, color=(100,100,100))

p2 = add_para(doc, space_after=1)
r3 = p2.add_run('Organizer, Logistics Team Lead and Ambassador   ')
set_font(r3, bold=True, size=9.5)
r4 = p2.add_run('Oct 2022 & Oct 2023')
set_font(r4, size=9.5, color=(100,100,100))

p3 = add_para(doc, space_after=4)
p3.paragraph_format.left_indent = Cm(0.5)
r5 = p3.add_run('• Organised NASA Space Apps Challenge at Salim Habib University (2022) and Habib University (2023); led logistics team ensuring smooth operations.')
set_font(r5, size=9.5)

# ── SKILLS ────────────────────────────────────────────────────
add_section_heading(doc, 'Skills & Language Proficiencies')

skills_data = [
    ('Technologies:', 'OpenCV, TensorFlow, Keras, Scikit-learn, PyTorch, MediaPipe, ARCore, Unity, LangChain, Hugging Face Transformers, Docker, Azure AI Services, Jupyter Notebook, Git, Flutter, Firebase, Arduino IDE'),
    ('Technical Skills:', 'AI, Machine Learning, Deep Learning, Computer Vision, NLP, Agentic AI & LLM Orchestration, Robotics, AR/VR, Model Training & Evaluation, Data Analysis, Feature Engineering'),
    ('Programming Languages:', 'Python, C++, Java, JavaFX, Dart, C#, LaTeX'),
    ('Soft Skills:', 'Communication, Leadership, Problem-Solving, Critical Thinking, Team Management, Research'),
    ('Languages:', 'Sindhi (Native), Urdu (Native), English (Fluent)'),
]

for label, value in skills_data:
    p = add_para(doc, space_before=2, space_after=1)
    r1 = p.add_run(label + '  ')
    set_font(r1, bold=True, size=9.5)
    r2 = p.add_run(value)
    set_font(r2, size=9.5)

# ── AWARDS ────────────────────────────────────────────────────
add_section_heading(doc, 'Awards')

awards = [
    ('Gold Medal', 'Presented by Department of Information Technology, Salim Habib University'),
    ('Galactic Problem Solver', 'Presented by NASA Space App Challenge (Oct 2022, Oct 2023)'),
    ('Dean Honor Award', 'Presented by Dean, Department of Information Technology, Salim Habib University (Spring 2021, Spring 2022, Fall 2022, Spring 2023, Fall 2023, Spring 2024, Fall 2024)'),
]

for title, detail in awards:
    p = add_para(doc, space_before=2, space_after=1)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(f'• {title}  ')
    set_font(r1, bold=True, size=9.5)
    r2 = p.add_run(f'– {detail}')
    set_font(r2, size=9.5)

# ── CERTIFICATIONS ────────────────────────────────────────────
add_section_heading(doc, 'Online Courses & Certifications')

certs = [
    ('AI Agents in LangGraph', 'DeepLearning.AI – Coursera', False),
    ('Supervised Machine Learning: Regression and Classification', 'DeepLearning.AI – Coursera', False),
    ('IBM RAG and Agentic AI', 'IBM – Coursera', True),
    ('IBM AI Developer', 'IBM – Coursera', True),
    ('Fundamentals of Agentic AI and DACA AI-First Development', 'Panaversity', True),
]

for title, issuer, ongoing in certs:
    p = add_para(doc, space_before=2, space_after=1)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(f'• {title}  ')
    set_font(r1, bold=True, size=9.5)
    r2 = p.add_run(f'– {issuer}')
    set_font(r2, size=9.5)
    if ongoing:
        r3 = p.add_run(' – ongoing')
        set_font(r3, size=9.5, color=(120,120,120))

# Save
out = '/data/data/com.termux/files/home/.openclaw/workspace/Muhammad_Fahad_Imdad_CV_Updated.docx'
doc.save(out)
print(f'Saved: {out}')
